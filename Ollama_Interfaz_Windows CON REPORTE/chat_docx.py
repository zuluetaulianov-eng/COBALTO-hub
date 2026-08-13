import os
from dataclasses import dataclass, field

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

from informe_osint import (
    BG_PAGE, BG_PANEL, BG_INPUT, BORDER,
    ACCENT, VERDE, ROJO, TXT, TXT_DIM, TXT_TITLE,
    FONT_MONO, FONT_UI,
    _set_cell_bg, _set_cell_border, _estilizar_celda, _no_partir_fila,
    _run, _tabla_fija, _fondo_pagina_visible, _pie_pagina,
)


@dataclass
class MensajeChat:
    role: str
    content: str


@dataclass
class ChatData:
    nombre_usuario: str = "Usuario"
    modelo: str = "llama3.2"
    temperatura: float = 0.7
    fecha: str = ""
    mensajes: list = field(default_factory=list)

    def to_dict(self) -> dict:
        base = {k: getattr(self, k) for k in self.__dataclass_fields__}
        base["mensajes"] = [m.__dict__ for m in self.mensajes]
        return base

    @classmethod
    def from_dict(cls, d: dict):
        d = dict(d)
        mensajes = d.get("mensajes", [])
        d["mensajes"] = [MensajeChat(**m) if isinstance(m, dict) else m
                         for m in mensajes]
        return cls(**{k: v for k, v in d.items() if k in cls.__dataclass_fields__})


def _caja_encabezado(doc, datos: ChatData):
    table = doc.add_table(rows=1, cols=1)
    _tabla_fija(table, [7.5])
    _no_partir_fila(table.rows[0])
    cell = table.cell(0, 0)
    _estilizar_celda(cell)
    p = cell.paragraphs[0]
    _run(p, "EL OJO DEL COPORO", FONT_MONO, 18, ACCENT, True)
    _linea(p, "[CONFIDENCIAL]", FONT_UI, 9, ROJO, True)
    _linea(p, "TRANSCRIPCIÓN DE ANÁLISIS DE IA", FONT_UI, 14, TXT_TITLE, True)
    doc.add_paragraph()


def _linea(p, texto, font=FONT_UI, size=9.5, color=TXT, bold=False):
    r = p.add_run()
    r.add_break()
    r = p.add_run(texto)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.bold = bold


def _tabla_meta(doc, datos: ChatData):
    table = doc.add_table(rows=2, cols=2)
    _tabla_fija(table, [3.5, 3.5])
    filas = [
        ("Usuario: ", f"{datos.nombre_usuario} | Modelo: {datos.modelo}"),
        ("Temperatura: ", f"{datos.temperatura:.2f}"),
    ]
    for i in range(2):
        cell = table.cell(0, i)
        _estilizar_celda(cell)
        p = cell.paragraphs[0]
        if i == 0:
            _run(p, filas[0][0], FONT_MONO, 8.5, TXT_DIM, True)
            _run(p, filas[0][1], FONT_MONO, 8.5, ACCENT)
        else:
            _run(p, "Fecha: ", FONT_MONO, 8.5, TXT_DIM, True)
            _run(p, datos.fecha or "N/D", FONT_MONO, 8.5, ACCENT)
    for i, (lbl, val) in enumerate([filas[1], ("Mensajes: ", f"{len(datos.mensajes)}")]):
        cell = table.cell(1, i)
        _estilizar_celda(cell)
        p = cell.paragraphs[0]
        _run(p, lbl, FONT_MONO, 8.5, TXT_DIM, True)
        _run(p, val, FONT_MONO, 8.5, ACCENT)
    doc.add_paragraph()


def _tarjeta_mensaje(doc, mensaje: MensajeChat):
    table = doc.add_table(rows=1, cols=1)
    _tabla_fija(table, [7.5])
    _no_partir_fila(table.rows[0])
    cell = table.cell(0, 0)
    _estilizar_celda(cell)
    _toggle_borde_tc(cell, VERDE)

    p = cell.paragraphs[0]
    _run(p, "OLLAMA  ▸  ANÁLISIS DE IA", FONT_MONO, 9, VERDE, True)
    p2 = cell.add_paragraph()
    _run(p2, mensaje.content, FONT_UI, 10, TXT)
    doc.add_paragraph()


def _toggle_borde_tc(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        return
    for edge in ("top", "left", "bottom", "right"):
        b = tc_borders.find(qn(f"w:{edge}"))
        if b is not None:
            b.set(qn("w:color"), color_hex)


def _tabla_estadisticas(doc, mensajes):
    n_user = sum(1 for m in mensajes if m.role == "user")
    n_bot = sum(1 for m in mensajes if m.role == "assistant")
    palabras = sum(len(m.content.split()) for m in mensajes)

    table = doc.add_table(rows=2, cols=3)
    _tabla_fija(table, [2.5, 2.5, 2.5])
    headers = ["Consultas", "Respuestas", "Palabras totales"]
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        _set_cell_border(cell)
        _set_cell_bg(cell, BG_INPUT)
        _run(cell.paragraphs[0], h, FONT_MONO, 8.5, ACCENT, True)
    for j, val in enumerate([str(n_user), str(n_bot), str(palabras)]):
        cell = table.cell(1, j)
        _estilizar_celda(cell)
        _run(cell.paragraphs[0], val, FONT_MONO, 8.5, TXT)
    doc.add_paragraph()


def generar_transcripcion(datos: ChatData,
                          output="transcripcion_ia.docx") -> str:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = FONT_UI
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(TXT)

    _fondo_pagina_visible(doc)
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    _pie_pagina(doc)
    _caja_encabezado(doc, datos)
    _tabla_meta(doc, datos)

    if datos.mensajes:
        p = doc.add_paragraph()
        _run(p, "ANÁLISIS DE LA IA", FONT_MONO, 11, ACCENT, True)
        for m in datos.mensajes:
            if m.role == "assistant":
                _tarjeta_mensaje(doc, m)

        p = doc.add_paragraph()
        _run(p, "RESUMEN DE LA SESIÓN", FONT_MONO, 11, ACCENT, True)
        _tabla_estadisticas(doc, datos.mensajes)
    else:
        p = doc.add_paragraph()
        _run(p, "Sin mensajes que exportar.", FONT_UI, 10, TXT_DIM)

    doc.save(output)
    return os.path.abspath(output)


def chat_desde_historial(historial, nombre_usuario, modelo,
                         temperatura, fecha=""):
    mensajes = [MensajeChat(role=m["role"], content=m["content"])
                for m in historial if m.get("content")]
    return ChatData(nombre_usuario=nombre_usuario or "Usuario",
                    modelo=modelo, temperatura=temperatura,
                    fecha=fecha, mensajes=mensajes)


if __name__ == "__main__":
    demo = ChatData(
        nombre_usuario="Ulianov",
        modelo="llama3.2",
        temperatura=0.7,
        fecha="12/08/2026 18:00",
        mensajes=[
            MensajeChat("user", "¿Qué es la entropía en termodinámica?"),
            MensajeChat("assistant",
                        "La entropía mide el grado de desorden o la cantidad de "
                        "energía no aprovechable de un sistema. En un sistema "
                        "cerrado, la entropía tiende a aumentar con el tiempo "
                        "(segunda ley de la termodinámica)."),
        ],
    )
    r = generar_transcripcion(demo)
    print(f"Transcripción generada: {r}")