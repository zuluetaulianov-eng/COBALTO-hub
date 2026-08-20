"""export_informe_osint.py - Generador de informe OSINT en DOCX (estilo cyber/dark).

Portado de 'Ollama_Interfaz_Windows CON REPORTE/informe_osint.py' al pipeline
Cobalto. Modelo de datos serializable (InformeData/Documento) + render con
python-docx y OXML para tarjetas, fondo oscuro, paginación y logo generado con Pillow.
"""

import io
import os
import tempfile
from dataclasses import dataclass, field
from datetime import datetime
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw

# ─── Paleta OSINT (Modo Imprimible / Legal Formal) ───────────────────────────
BG_PAGE = "FFFFFF"
BG_PANEL = "F8FAFC"
BG_INPUT = "F1F5F9"
BORDER = "CBD5E1"
ACCENT = "0284C7"
ACCENT_SOF = "0369A1"
VERDE = "16A34A"
VERDE_DO = "15803D"
ROJO = "DC2626"
TXT = "1E293B"
TXT_DIM = "475569"
TXT_TITLE = "0F172A"

FONT_MONO = "Courier New"
FONT_UI = "Segoe UI"


@dataclass
class Documento:
    doc_num: str
    titulo: str
    fuente: str
    score_sentimiento: float = 0.0
    url: str = ""
    analisis: str = ""
    contenido: str = ""

    def to_dict(self) -> dict:
        return {k: getattr(self, k) for k in self.__dataclass_fields__}

    @classmethod
    def from_dict(cls, d: dict):
        d = dict(d)
        d = {k: v for k, v in d.items() if k in cls.__dataclass_fields__}
        d.setdefault("score_sentimiento", 0.0)
        return cls(**d)


@dataclass
class InformeData:
    codigo: str
    fecha_creacion: str
    autor: str
    institucion: str
    fuente_datos: str
    fecha_analisis: str
    titulo_seccion: str = "ANÁLISIS POR DOCUMENTO"
    documentos: list = field(default_factory=list)
    total_analizados: int = 0
    doc_con_bot: int = 0
    nivel_alerta: str = "MONITOREO NORMAL"
    niveles_bot: list = field(default_factory=list)
    fuentes_bot: list = field(default_factory=list)

    def to_dict(self) -> dict:
        base = {k: getattr(self, k) for k in self.__dataclass_fields__}
        base["documentos"] = [d.to_dict() if isinstance(d, Documento) else d for d in self.documentos]
        return base

    @classmethod
    def from_dict(cls, d: dict):
        d = dict(d)
        docs = d.get("documentos", [])
        d["documentos"] = [Documento.from_dict(x) if isinstance(x, dict) else x for x in docs]
        d = {k: v for k, v in d.items() if k in cls.__dataclass_fields__}
        return cls(**d)


# ─── Helpers OXML ─────────────────────────────────────────────────────────────
def _set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cell_border(cell, color_hex=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex)
        tc_borders.append(b)
    tc_pr.insert(0, tc_borders)


def _estilizar_celda(cell, bg=BG_PANEL):
    _set_cell_border(cell)
    _set_cell_bg(cell, bg)


def _no_partir_fila(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def _run(p, texto, font=FONT_UI, size=9.5, color=TXT, bold=False, italic=False):
    r = p.add_run(texto)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.bold = bold
    r.italic = italic
    return r


def _linea(p, texto, font=FONT_UI, size=9.5, color=TXT, bold=False, italic=False):
    r = _run(p, "", font, size, color, bold, italic)
    r.add_break()
    r = p.add_run(texto)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.bold = bold
    r.italic = italic
    return r


def _tabla_fija(table, anchos):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    for row in table.rows:
        for idx, w in enumerate(anchos):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(w)


def _fondo_pagina_visible(doc, color_hex=BG_PAGE):
    if color_hex and color_hex.upper() != "FFFFFF":
        doc.element.insert(0, OxmlElement("w:background"))
        doc.element[0].set(qn("w:color"), color_hex)
        doc.settings.element.append(OxmlElement("w:displayBackgroundShape"))


def _pie_pagina(doc):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "EL OJO DEL COPORO - OSINT CONFIDENCIAL  |  Página ", FONT_MONO, 8, ACCENT)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r = p.add_run()
    r.font.name = FONT_MONO
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(ACCENT)
    r._r.append(fld1)
    r._r.append(instr)
    r._r.append(fld2)


def _crear_logo_temporal():
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    tmp.close()
    size = 240
    img = Image.new("RGB", (size, size), "#" + BG_PANEL)
    d = ImageDraw.Draw(img)
    s = size / 100.0
    d.polygon(
        [(50 * s, 12 * s), (72 * s, 28 * s), (78 * s, 48 * s), (66 * s, 78 * s),
         (50 * s, 85 * s), (34 * s, 78 * s), (22 * s, 48 * s), (28 * s, 28 * s)],
        outline="#1f6feb", width=3)
    d.polygon(
        [(40 * s, 22 * s), (52 * s, 18 * s), (46 * s, 30 * s)], fill="#0d1117", outline="#58a6ff")
    d.ellipse([36 * s, 32 * s, 64 * s, 60 * s], outline="#58a6ff", width=3)
    d.ellipse([45 * s, 41 * s, 55 * s, 51 * s], fill="#58a6ff")
    d.line([(50 * s, 35 * s), (50 * s, 41 * s)], fill="#388bfd", width=3)
    d.line([(50 * s, 51 * s), (50 * s, 75 * s)], fill="#388bfd", width=2)
    d.line([(27 * s, 46 * s), (36 * s, 46 * s)], fill="#388bfd", width=2)
    d.line([(64 * s, 46 * s), (73 * s, 46 * s)], fill="#388bfd", width=2)
    d.ellipse([30 * s, 44 * s, 70 * s, 48 * s], outline="#30363d", width=1)
    img.save(tmp.name)
    return tmp.name


def _encabezado(doc, logo_path):
    table = doc.add_table(rows=1, cols=2)
    _tabla_fija(table, [1.1, 5.4])
    logo_cell = table.cell(0, 0)
    _estilizar_celda(logo_cell)
    p = logo_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(logo_path, width=Inches(0.85))

    title_cell = table.cell(0, 1)
    _estilizar_celda(title_cell)
    p = title_cell.paragraphs[0]
    _run(p, "EL OJO DEL COPORO", FONT_MONO, 16, ACCENT, True)
    _linea(p, "[CONFIDENCIAL]", FONT_UI, 8.5, ROJO, True)
    _linea(p, "INFORME DE INTELIGENCIA OSINT", FONT_UI, 12, TXT_TITLE, True)
    doc.add_paragraph()


def _meta(doc, d: InformeData):
    table = doc.add_table(rows=3, cols=2)
    _tabla_fija(table, [3.25, 3.25])
    filas = [
        [("Código: ", d.codigo), ("Fecha de Creación: ", d.fecha_creacion)],
        [("Autor: ", d.autor), ("Institución: ", d.institucion)],
    ]
    for i, (label, value) in enumerate(filas[0]):
        cell = table.cell(0, i)
        _estilizar_celda(cell)
        p = cell.paragraphs[0]
        _run(p, label, FONT_MONO, 8.5, TXT_DIM, True)
        _run(p, value, FONT_MONO, 8.5, ACCENT)
    for i, (label, value) in enumerate(filas[1]):
        cell = table.cell(1, i)
        _estilizar_celda(cell)
        p = cell.paragraphs[0]
        _run(p, label, FONT_MONO, 8.5, TXT_DIM, True)
        _run(p, value, FONT_MONO, 8.5, ACCENT)
    merged = table.cell(2, 0).merge(table.cell(2, 1))
    _estilizar_celda(merged)
    p = merged.paragraphs[0]
    _run(p, "Fuente de Datos: ", FONT_MONO, 8.5, TXT_DIM, True)
    _run(p, d.fuente_datos, FONT_MONO, 8.5, ACCENT)
    doc.add_paragraph()


def _seccion(doc, titulo, subtitulo=""):
    p = doc.add_paragraph()
    _run(p, titulo, FONT_MONO, 10.5, ACCENT, True)
    if subtitulo:
        p2 = doc.add_paragraph()
        _run(p2, subtitulo, FONT_UI, 9, TXT_DIM)


def _caja_intel(cell, d: Documento):
    tbl = cell.add_table(rows=1, cols=1)
    _tabla_fija(tbl, [6.1])
    c = tbl.cell(0, 0)
    _set_cell_border(c, VERDE_DO)
    _set_cell_bg(c, BG_PANEL)
    p = c.paragraphs[0]
    _run(p, "ANÁLISIS DE INTELIGENCIA", FONT_MONO, 9, VERDE, True)
    p2 = c.add_paragraph()
    _run(p2, d.analisis, FONT_UI, 9.5, TXT)


def _tarjeta_doc(doc, d: Documento):
    table = doc.add_table(rows=1, cols=1)
    _tabla_fija(table, [6.5])
    _no_partir_fila(table.rows[0])
    cell = table.cell(0, 0)
    _estilizar_celda(cell)

    p = cell.paragraphs[0]
    _run(p, f"📇 [TARJETA NOTICIOSA #{d.doc_num}] ", FONT_MONO, 9, ACCENT_SOF, True)
    _run(p, d.titulo, FONT_UI, 11, TXT_TITLE, True)

    p2 = cell.add_paragraph()
    _run(p2, f"📡 Fuente: {d.fuente} ", FONT_MONO, 8.5, TXT_DIM, True)
    _run(p2, f"| ⚖️ Score Sentimiento: {d.score_sentimiento:.2f} ", FONT_MONO, 8.5, TXT_DIM, True)
    if d.url:
        _run(p2, f"| 🔗 URL: {d.url}", FONT_MONO, 8.5, ACCENT)

    if d.analisis:
        _caja_intel(cell, d)

    p3 = cell.add_paragraph()
    _run(p3, "📝 Resumen / Contenido Noticioso: " + d.contenido, FONT_UI, 9, TXT, italic=False)
    doc.add_paragraph()


def _resumen_bots(doc, d: InformeData):
    pct = d.doc_con_bot / d.total_analizados * 100 if d.total_analizados else 0.0
    table = doc.add_table(rows=1, cols=1)
    _tabla_fija(table, [6.5])
    _no_partir_fila(table.rows[0])
    cell = table.cell(0, 0)
    _estilizar_celda(cell)
    p = cell.paragraphs[0]
    _run(p, f"De un total de {d.total_analizados} documentos analizados, "
            f"{d.doc_con_bot} presentan indicios de actividad automatizada "
            f"({pct:.1f}%). Nivel de alerta actual: ", FONT_UI, 9.5, TXT)
    r = p.add_run(f"[{d.nivel_alerta}]")
    r.font.name = FONT_MONO
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(ACCENT)
    r.bold = True


def _tabla_datos(doc, titulo, headers, filas, anchos, columna_negrita=0):
    p = doc.add_paragraph()
    _run(p, titulo, FONT_MONO, 9, TXT_TITLE, True)

    table = doc.add_table(rows=1, cols=len(headers))
    _tabla_fija(table, anchos)
    _no_partir_fila(table.rows[0])
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        _set_cell_border(cell)
        _set_cell_bg(cell, BG_INPUT)
        pc = cell.paragraphs[0]
        _run(pc, h, FONT_MONO, 8.5, ACCENT, True)

    for fila in filas:
        row = table.add_row()
        _no_partir_fila(row)
        for j, valor in enumerate(fila):
            cell = row.cells[j]
            _estilizar_celda(cell)
            pc = cell.paragraphs[0]
            _run(pc, str(valor), FONT_MONO, 8.5, TXT, bold=(j == columna_negrita))


def generar_informe_osint(datos: InformeData, output="informe_inteligencia_coporo.docx") -> str:
    logo_path = None
    try:
        logo_path = _crear_logo_temporal()
        doc = Document()

        normal = doc.styles["Normal"]
        normal.font.name = FONT_UI
        normal.font.size = Pt(9.5)
        normal.font.color.rgb = RGBColor.from_string(TXT)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(2)

        _fondo_pagina_visible(doc)
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        _pie_pagina(doc)
        _encabezado(doc, logo_path)
        _meta(doc, datos)
        _seccion(doc, datos.titulo_seccion,
                 f"Documentos analizados: {len(datos.documentos)} | Fecha: {datos.fecha_analisis}")
        for d in datos.documentos:
            _tarjeta_doc(doc, d)

        _seccion(doc, "ANÁLISIS DE ACTIVIDAD DE BOTS")
        _resumen_bots(doc, datos)
        if datos.niveles_bot:
            _tabla_datos(
                doc,
                "Distribución por Nivel de Bot Score",
                ["Nivel", "Documentos", "% del Total"],
                datos.niveles_bot,
                [3.0, 1.75, 1.75],
            )
        if datos.fuentes_bot:
            _tabla_datos(
                doc,
                "Fuentes con Mayor Actividad de Bot",
                ["#", "Fuente", "Documentos con Bot", "Bot Score Promedio"],
                datos.fuentes_bot,
                [0.6, 2.0, 1.95, 1.95],
            )

        doc.save(output)
        return os.path.abspath(output)
    finally:
        if logo_path and os.path.exists(logo_path):
            os.unlink(logo_path)


def generar_informe_osint_bytes(datos: InformeData) -> bytes:
    """Genera el documento y retorna los bytes (para servir por API sin dejar archivos)."""
    buf = io.BytesIO()
    logo_path = None
    try:
        logo_path = _crear_logo_temporal()
        doc = Document()

        normal = doc.styles["Normal"]
        normal.font.name = FONT_UI
        normal.font.size = Pt(9.5)
        normal.font.color.rgb = RGBColor.from_string(TXT)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(2)

        _fondo_pagina_visible(doc)
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        _pie_pagina(doc)
        _encabezado(doc, logo_path)
        _meta(doc, datos)
        _seccion(doc, datos.titulo_seccion,
                 f"Documentos analizados: {len(datos.documentos)} | Fecha: {datos.fecha_analisis}")
        for d in datos.documentos:
            _tarjeta_doc(doc, d)

        _seccion(doc, "ANÁLISIS DE ACTIVIDAD DE BOTS")
        _resumen_bots(doc, datos)
        if datos.niveles_bot:
            _tabla_datos(
                doc,
                "Distribución por Nivel de Bot Score",
                ["Nivel", "Documentos", "% del Total"],
                datos.niveles_bot,
                [3.0, 1.75, 1.75],
            )
        if datos.fuentes_bot:
            _tabla_datos(
                doc,
                "Fuentes con Mayor Actividad de Bot",
                ["#", "Fuente", "Documentos con Bot", "Bot Score Promedio"],
                datos.fuentes_bot,
                [0.6, 2.0, 1.95, 1.95],
            )

        doc.save(buf)
        return buf.getvalue()
    finally:
        if logo_path and os.path.exists(logo_path):
            os.unlink(logo_path)


def _to_doc_num(idx: int) -> str:
    return str(idx + 1)


def _a_float(valor, por_defecto=0.0):
    try:
        return float(valor)
    except (TypeError, ValueError):
        return por_defecto


def _a_texto(valor) -> str:
    if valor is None:
        return ""
    return str(valor)


def build_informe_desde_entries(entries: List[dict], max_docs: int = 20,
                                analisis_por_entry: dict = None) -> InformeData:
    """Construye InformeData a partir de las entradas del contexto Cobalto.

    Cada entrada aporta título, fuente, resumen/contenido y, si existe, el análisis
    de IA (actores/amenaza). Los scores de bot y sentimiento se toman de los campos
    del dashboard (bot_probability, sentiment_score) con mapeo defensivo.
    """
    ahora = datetime.now()
    info = InformeData(
        codigo=f"INT-OSINT-{ahora.strftime('%Y%m%d')}-{ahora.strftime('%H%M')}",
        fecha_creacion=ahora.strftime("%d de %B de %Y, %H:%M"),
        autor="Sistema COBALTO HUB",
        institucion="EL OJO DEL COPORO",
        fuente_datos=f"COBALTO HUB Dashboard ({len(entries)} entradas en contexto)",
        fecha_analisis=ahora.strftime("%d/%m/%Y %H:%M"),
        titulo_seccion="ANÁLISIS POR DOCUMENTO",
    )
    analisis_por_entry = analisis_por_entry or {}

    for idx, entry in enumerate(entries[:max_docs]):
        eid = str(entry.get("id", entry.get("title", "")))
        analisis = ""
        if eid in analisis_por_entry:
            a = analisis_por_entry.get(eid, {})
            if isinstance(a, dict):
                piezas = [a.get("analisis", "")]
                actores = a.get("actores")
                if isinstance(actores, list) and actores and actores != ["N/A"]:
                    piezas.append("Actores: " + ", ".join(actores))
                amenaza = a.get("amenaza")
                if amenaza and amenaza not in ("Desconocida", "N/A"):
                    piezas.append(f"Threat Level: {amenaza}")
                analisis = " | ".join([x for x in piezas if x])
        if not analisis:
            analisis = _a_texto(entry.get("analysis", entry.get("intro", entry.get("resumen", ""))))

        doc = Documento(
            doc_num=_to_doc_num(idx),
            titulo=_a_texto(entry.get("title", entry.get("titulo", "Sin título"))),
            fuente=_a_texto(entry.get("source", entry.get("fuente", "desconocido"))),
            score_sentimiento=_a_float(entry.get("sentiment_score", 0.0)),
            url=_a_texto(entry.get("link", entry.get("url", ""))),
            analisis=analisis,
            contenido=_a_texto(entry.get("summary", entry.get("texto", ""))),
        )
        info.documentos.append(doc)

    info.total_analizados = max(len(entries), info.total_analizados)
    info.doc_con_bot = sum(
        1 for e in entries[:max_docs] if _a_float(e.get("bot_probability", 0.0)) > 0.5
    )
    if info.doc_con_bot:
        info.nivel_alerta = "POSIBLE CAMPAÑA COORDINADA"
    return info


def datos_ejemplo() -> InformeData:
    return InformeData(
        codigo="INT-OSINT-2026-0030",
        fecha_creacion="16 de junio de 2026, 17:46 UTC",
        autor="Analista COPORO",
        institucion="EL OJO DEL COPORO",
        fuente_datos="EL OJO DEL COPORO v10.0.0",
        fecha_analisis="16/06/2026 17:46",
        documentos=[
            Documento(
                doc_num="1",
                titulo="Ejemplo de documento analizado",
                fuente="rss",
                url="https://example.com",
                analisis="Análisis de inteligencia de ejemplo para validar el render DOCX.",
                contenido="Contenido de ejemplo.",
            ),
        ],
        total_analizados=1,
        doc_con_bot=0,
        niveles_bot=[],
        fuentes_bot=[],
    )


if __name__ == "__main__":
    ruta = generar_informe_osint(datos_ejemplo())
    print(f"Informe generado exitosamente: {ruta}")
