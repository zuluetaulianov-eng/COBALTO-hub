"""
EXPORT SITREP DOCX (Generador Determinista de Informes Word / DOCX)
===================================================================
Construye informes de situación (SITREP) en formato Microsoft Word (.docx)
con diseño táctico, tablas estilizadas, metadata de sistema y formato formal.
SIN DEPENDENCIAS DE IA.
"""

import io
import json
import logging
import os
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger("SitrepDocxNoIA")

_INVALID_XML_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _sanitizar(texto: str) -> str:
    return _INVALID_XML_RE.sub("", str(texto or ""))


def _set_cell_bg(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cell_border(cell, color_hex: str = "CBD5E1"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex)
        tc_borders.append(b)
    tc_pr.insert(0, tc_borders)


def generate_sitrep_docx_bytes(contexto: dict) -> bytes:
    """Genera un informe SITREP en formato Word (.docx) como stream de bytes."""
    doc = Document()

    # Estilos globales
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor(30, 41, 59)

    # Márgenes de página
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Header / Banner Táctico
    table_header = doc.add_table(rows=1, cols=1)
    cell_h = table_header.cell(0, 0)
    _set_cell_bg(cell_h, "0F172A")
    p_h = cell_h.paragraphs[0]
    p_h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run_title = p_h.add_run("  SITREP OPERACIONAL - REPORTE DE SITUACIÓN\n")
    run_title.font.name = "Courier New"
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = RGBColor(0, 229, 255)
    run_title.bold = True

    run_sub = p_h.add_run(f"  CLASIFICACIÓN: NO CLASIFICADO | GENERADO: {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}")
    run_sub.font.name = "Segoe UI"
    run_sub.font.size = Pt(8.5)
    run_sub.font.color.rgb = RGBColor(148, 163, 184)

    doc.add_paragraph()

    # Metadata de Resumen
    p_meta = doc.add_paragraph()
    r_meta_t = p_meta.add_run("1. RESUMEN EJECUTIVO Y METADATOS\n")
    r_meta_t.font.name = "Courier New"
    r_meta_t.font.size = Pt(11)
    r_meta_t.font.color.rgb = RGBColor(2, 132, 199)
    r_meta_t.bold = True

    entries = contexto.get("entries", contexto.get("all_entries", []))
    alerts = contexto.get("alerts", [])

    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.cell(0, 0).paragraphs[0].text = f"Total Entradas Analizadas: {len(entries)}"
    meta_table.cell(0, 1).paragraphs[0].text = f"Alertas Activas: {len(alerts)}"
    meta_table.cell(1, 0).paragraphs[0].text = f"Estado del Sistema: OPERATIVO (100% Determinista)"
    meta_table.cell(1, 1).paragraphs[0].text = f"Modo de Generación: Reglas Heurísticas (Sin IA)"

    for row in meta_table.rows:
        for cell in row.cells:
            _set_cell_bg(cell, "F8FAFC")
            _set_cell_border(cell, "CBD5E1")

    doc.add_paragraph()

    # Tabla de Entradas Noticiosas
    p_entries = doc.add_paragraph()
    r_ent_t = p_entries.add_run("2. REGISTRO DE EVENTOS NOTICIOSOS ANALIZADOS\n")
    r_ent_t.font.name = "Courier New"
    r_ent_t.font.size = Pt(11)
    r_ent_t.font.color.rgb = RGBColor(2, 132, 199)
    r_ent_t.bold = True

    if entries:
        t_data = doc.add_table(rows=1, cols=4)
        hdr_cells = t_data.rows[0].cells
        hdr_titles = ["#", "Título / Evento", "Fuente", "Nivel Amenaza"]
        for idx, text in enumerate(hdr_titles):
            hdr_cells[idx].paragraphs[0].text = text
            _set_cell_bg(hdr_cells[idx], "F1F5F9")
            _set_cell_border(hdr_cells[idx], "94A3B8")

        for i, entry in enumerate(entries[:25]):
            row_cells = t_data.add_row().cells
            row_cells[0].paragraphs[0].text = str(i + 1)
            row_cells[1].paragraphs[0].text = _sanitizar(entry.get("title", entry.get("titulo", "Sin título")))[:60]
            row_cells[2].paragraphs[0].text = _sanitizar(entry.get("source", entry.get("fuente", "N/A")))
            
            analisis = entry.get("analisis_determinista", {})
            amenaza = analisis.get("nivel_amenaza", "MEDIA")
            row_cells[3].paragraphs[0].text = amenaza
            _set_cell_bg(row_cells[3], "FEF2F2" if amenaza in ("CRÍTICA", "ALTA") else "F8FAFC")

            for c in row_cells:
                _set_cell_border(c, "CBD5E1")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
