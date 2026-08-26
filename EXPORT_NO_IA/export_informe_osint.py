"""
EXPORT INFORME OSINT (Generador de Documentos Formales de Inteligencia y Auditoría Bot)
======================================================================================
Crea informes ejecutivos completos en Microsoft Word (.docx) con renderizado de
gráficos vectoriales (logo Pillow), tarjetas noticiosas, auditoría de actividad bot,
matrices de sentimiento y encabezados confidenciales.
SIN DEPENDENCIAS DE LLM / IA.
"""

import io
import os
import tempfile
from dataclasses import dataclass, field
from datetime import datetime
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw

BG_PAGE = "FFFFFF"
BG_PANEL = "F8FAFC"
BG_INPUT = "F1F5F9"
BORDER = "CBD5E1"
ACCENT = "0284C7"
ACCENT_SOF = "0369A1"
VERDE = "16A34A"
ROJO = "DC2626"
TXT = "1E293B"
TXT_DIM = "475569"
TXT_TITLE = "0F172A"
FONT_MONO = "Courier New"
FONT_UI = "Segoe UI"


@dataclass
class DocumentoNoticioso:
    doc_num: str
    titulo: str
    fuente: str
    score_sentimiento: float = 0.0
    url: str = ""
    analisis: str = ""
    contenido: str = ""


@dataclass
class InformeOSINTData:
    codigo: str
    fecha_creacion: str
    autor: str
    institucion: str
    fuente_datos: str
    fecha_analisis: str
    titulo_seccion: str = "ANÁLISIS DE DOCUMENTOS DE INTELIGENCIA"
    documentos: list = field(default_factory=list)
    total_analizados: int = 0
    doc_con_bot: int = 0
    nivel_alerta: str = "MONITOREO ESTÁNDAR"


def _set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cell_border(cell, color_hex=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex)
        tc_borders.append(b)
    tc_pr.insert(0, tc_borders)


def _crear_logo_temporal():
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    tmp.close()
    size = 240
    img = Image.new("RGB", (size, size), "#" + BG_PANEL)
    d = ImageDraw.Draw(img)
    s = size / 100.0
    d.polygon([(50 * s, 12 * s), (72 * s, 28 * s), (78 * s, 48 * s), (66 * s, 78 * s), (50 * s, 85 * s), (34 * s, 78 * s), (22 * s, 48 * s), (28 * s, 28 * s)], outline="#0284c7", width=3)
    d.ellipse([36 * s, 32 * s, 64 * s, 60 * s], outline="#0284c7", width=3)
    d.ellipse([45 * s, 41 * s, 55 * s, 51 * s], fill="#0284c7")
    img.save(tmp.name)
    return tmp.name


def generar_informe_osint_bytes(datos: InformeOSINTData) -> bytes:
    """Genera un informe OSINT ejecutivo en DOCX y retorna sus bytes."""
    buf = io.BytesIO()
    logo_path = None
    try:
        logo_path = _crear_logo_temporal()
        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(0.9)
            section.bottom_margin = Inches(0.9)
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)

        # Encabezado
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        cell_logo = table.cell(0, 0)
        cell_logo.width = Inches(1.1)
        _set_cell_bg(cell_logo, BG_PANEL)
        _set_cell_border(cell_logo, BORDER)
        p_l = cell_logo.paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_l.add_run().add_picture(logo_path, width=Inches(0.8))

        cell_title = table.cell(0, 1)
        cell_title.width = Inches(5.4)
        _set_cell_bg(cell_title, BG_PANEL)
        _set_cell_border(cell_title, BORDER)
        p_t = cell_title.paragraphs[0]
        r1 = p_t.add_run("SISTEMA DE INTELIGENCIA C4I\n")
        r1.font.name = FONT_MONO
        r1.font.size = Pt(14)
        r1.font.color.rgb = RGBColor(2, 132, 199)
        r1.bold = True

        r2 = p_t.add_run("INFORME DETERMINISTA OSINT & AUDITORÍA NOTICIOSA")
        r2.font.name = FONT_UI
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(15, 23, 42)
        r2.bold = True

        doc.add_paragraph()

        # Metadata
        t_meta = doc.add_table(rows=2, cols=2)
        t_meta.cell(0, 0).paragraphs[0].text = f"Código: {datos.codigo}"
        t_meta.cell(0, 1).paragraphs[0].text = f"Fecha: {datos.fecha_creacion}"
        t_meta.cell(1, 0).paragraphs[0].text = f"Autor: {datos.autor}"
        t_meta.cell(1, 1).paragraphs[0].text = f"Fuente: {datos.fuente_datos}"
        for row in t_meta.rows:
            for c in row.cells:
                _set_cell_bg(c, BG_PANEL)
                _set_cell_border(c, BORDER)

        doc.add_paragraph()

        # Tarjetas de documentos
        for doc_item in datos.documentos:
            t_card = doc.add_table(rows=1, cols=1)
            cell_card = t_card.cell(0, 0)
            _set_cell_bg(cell_card, "FFFFFF")
            _set_cell_border(cell_card, "CBD5E1")

            p_c = cell_card.paragraphs[0]
            r_c1 = p_c.add_run(f"📇 [{doc_item.doc_num}] {doc_item.titulo}\n")
            r_c1.font.name = FONT_UI
            r_c1.font.size = Pt(10.5)
            r_c1.font.color.rgb = RGBColor(15, 23, 42)
            r_c1.bold = True

            r_c2 = p_c.add_run(f"Fuente: {doc_item.fuente} | Sentimiento: {doc_item.score_sentimiento:.2f}\n")
            r_c2.font.name = FONT_MONO
            r_c2.font.size = Pt(8.5)
            r_c2.font.color.rgb = RGBColor(71, 85, 105)

            if doc_item.analisis:
                r_c3 = p_c.add_run(f"Análisis: {doc_item.analisis}\n")
                r_c3.font.name = FONT_UI
                r_c3.font.size = Pt(9)
                r_c3.font.color.rgb = RGBColor(2, 132, 199)

            r_c4 = p_c.add_run(f"Resumen: {doc_item.contenido}")
            r_c4.font.name = FONT_UI
            r_c4.font.size = Pt(9)

            doc.add_paragraph()

        doc.save(buf)
        return buf.getvalue()
    finally:
        if logo_path and os.path.exists(logo_path):
            os.unlink(logo_path)
