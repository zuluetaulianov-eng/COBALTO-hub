"""Genera la plantilla DOCX template_sitrep.docx para el exportador SitRep."""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "..", "templates")
OUTPUT_PATH = os.path.join(TEMPLATE_DIR, "template_sitrep.docx")


def _set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def _add_styled_paragraph(doc, text, size=11, bold=False, color=None, alignment=None, font_name="Arial"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment:
        p.alignment = alignment
    return p


def generate_template():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    _add_styled_paragraph(
        doc, "SITREP COBALTO - INFORME DE INTELIGENCIA OSINT",
        size=18, bold=True, color=(0, 180, 255),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_styled_paragraph(
        doc, "Situation Report — Plataforma de Inteligencia C4I",
        size=10, bold=False, color=(100, 100, 100),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph()

    _add_styled_paragraph(
        doc, "CLASIFICACION: NO CLASIFICADO — USO INTERNO",
        size=9, bold=True, color=(200, 50, 50),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph()

    hr = doc.add_paragraph()
    hr_run = hr.add_run("─" * 80)
    hr_run.font.size = Pt(6)
    hr_run.font.color.rgb = RGBColor(0, 180, 255)

    _add_styled_paragraph(
        doc, "1. IDENTIFICACION DEL REPORTE",
        size=12, bold=True, color=(0, 180, 255),
    )

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Table Grid"
    meta_fields = [
        ("Version del SitRep:", "{{ sitrep_version }}"),
        ("Generado el:", "{{ generated_at }}"),
        ("Estado del Sistema:", "{{ system_status }}"),
        ("Ciclo ID:", "{{ cycle_id }}"),
    ]
    for i, (label, value) in enumerate(meta_fields):
        c0 = meta_table.rows[i].cells[0]
        c1 = meta_table.rows[i].cells[1]
        c0.text = label
        c1.text = value
        for c in (c0, c1):
            for p in c.paragraphs:
                p.style.font.size = Pt(10)
        _set_cell_shading(c0, "1A1A2E")

    doc.add_paragraph()

    _add_styled_paragraph(
        doc, "2. RESUMEN EJECUTIVO",
        size=12, bold=True, color=(0, 180, 255),
    )
    _add_styled_paragraph(doc, "{{ briefing_resumen }}", size=10)
    doc.add_paragraph()

    _add_styled_paragraph(
        doc, "3. METRICAS DEL SISTEMA",
        size=12, bold=True, color=(0, 180, 255),
    )

    metrics_table = doc.add_table(rows=6, cols=2)
    metrics_table.style = "Table Grid"
    metrics_fields = [
        ("Total Entradas:", "{{ total_entries }}"),
        ("Total Alertas:", "{{ total_alerts }}"),
        ("Circuit Breakers Abiertos:", "{{ cb_count }}"),
        ("Fuentes Totales:", "{{ total_sources }}"),
        ("Groq Disponible:", "{{ groq_available }}"),
        ("Nivel de Stress:", "{{ stress_level }}"),
    ]
    for i, (label, value) in enumerate(metrics_fields):
        c0 = metrics_table.rows[i].cells[0]
        c1 = metrics_table.rows[i].cells[1]
        c0.text = label
        c1.text = value
        _set_cell_shading(c0, "1A1A2E")

    doc.add_paragraph()

    if_has_alerts = doc.add_paragraph()
    if_has_alerts_run = if_has_alerts.add_run("{% if alerts %}")
    if_has_alerts_run.font.size = Pt(1)
    if_has_alerts_run.font.color.rgb = RGBColor(255, 255, 255)

    _add_styled_paragraph(
        doc, "4. ALERTAS ACTIVAS",
        size=12, bold=True, color=(255, 50, 50),
    )

    alert_header = doc.add_table(rows=1, cols=4)
    alert_header.style = "Table Grid"
    for ci, hdr in enumerate(["Tipo", "Severidad", "Fuente", "Timestamp"]):
        alert_header.rows[0].cells[ci].text = hdr
        _set_cell_shading(alert_header.rows[0].cells[ci], "330000")

    _add_styled_paragraph(doc, "{% for alert in alerts %}", size=1, color=(255, 255, 255))
    alert_row = doc.add_table(rows=2, cols=4)
    alert_row.style = "Table Grid"
    for ci, field in enumerate(["{{ alert.type }}", "{{ alert.severity }}", "{{ alert.source }}", "{{ alert.timestamp }}"]):
        alert_row.rows[0].cells[ci].text = field
    for ci in range(4):
        alert_row.rows[1].cells[ci].merge(alert_row.rows[1].cells[ci])
    alert_row.rows[1].cells[0].text = "{{ alert.title }}"
    _add_styled_paragraph(doc, "{% endfor %}", size=1, color=(255, 255, 255))

    endif_alerts = doc.add_paragraph()
    endif_alerts_run = endif_alerts.add_run("{% endif %}")
    endif_alerts_run.font.size = Pt(1)
    endif_alerts_run.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    if_outages = doc.add_paragraph()
    if_outages_run = if_outages.add_run("{% if outages %}")
    if_outages_run.font.size = Pt(1)
    if_outages_run.font.color.rgb = RGBColor(255, 255, 255)

    _add_styled_paragraph(
        doc, "5. APAGONES DE RED ACTIVOS",
        size=12, bold=True, color=(255, 150, 0),
    )

    outage_header = doc.add_table(rows=1, cols=3)
    outage_header.style = "Table Grid"
    for ci, hdr in enumerate(["ASN", "Pais", "Caida %"]):
        outage_header.rows[0].cells[ci].text = hdr
        _set_cell_shading(outage_header.rows[0].cells[ci], "332200")

    _add_styled_paragraph(doc, "{% for outage in outages %}", size=1, color=(255, 255, 255))
    outage_row = doc.add_table(rows=1, cols=3)
    outage_row.style = "Table Grid"
    for ci, field in enumerate(["{{ outage.asn }}", "{{ outage.country }}", "{{ outage.drop_percent }}"]):
        outage_row.rows[0].cells[ci].text = field
    _add_styled_paragraph(doc, "{% endfor %}", size=1, color=(255, 255, 255))

    endif_outages = doc.add_paragraph()
    endif_outages_run = endif_outages.add_run("{% endif %}")
    endif_outages_run.font.size = Pt(1)
    endif_outages_run.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    if_entries = doc.add_paragraph()
    if_entries_run = if_entries.add_run("{% if entries %}")
    if_entries_run.font.size = Pt(1)
    if_entries_run.font.color.rgb = RGBColor(255, 255, 255)

    _add_styled_paragraph(
        doc, "6. ENTRADAS OSINT RECIENTES",
        size=12, bold=True, color=(0, 180, 255),
    )

    entry_header = doc.add_table(rows=1, cols=5)
    entry_header.style = "Table Grid"
    for ci, hdr in enumerate(["#", "Titulo", "Fuente", "Fecha", "Critico"]):
        entry_header.rows[0].cells[ci].text = hdr
        _set_cell_shading(entry_header.rows[0].cells[ci], "1A2E1A")

    _add_styled_paragraph(doc, "{% for entry in entries %}", size=1, color=(255, 255, 255))
    entry_row = doc.add_table(rows=3, cols=5)
    entry_row.style = "Table Grid"
    for ci, field in enumerate([
        "{{ entry.idx }}", "{{ entry.title }}", "{{ entry.source }}",
        "{{ entry.published }}", "{{ entry.is_crisis }}"
    ]):
        entry_row.rows[0].cells[ci].text = field
    entry_row.rows[1].cells[0].merge(entry_row.rows[1].cells[4])
    entry_row.rows[1].cells[0].text = "Link: {{ entry.link }}"
    entry_row.rows[2].cells[0].merge(entry_row.rows[2].cells[4])
    entry_row.rows[2].cells[0].text = "{{ entry.summary }}"

    if_entry_analysis = entry_row.add_row()
    if_entry_analysis_cell = if_entry_analysis.cells[0]
    if_entry_analysis_cell.merge(if_entry_analysis.cells[4])
    if_entry_analysis_cell.text = ""
    ia_run = if_entry_analysis_cell.paragraphs[0].add_run(
        "{% if entry.analysis %}Analisis IA — Actores: {{ entry.analysis.actores }} | "
        "Amenaza: {{ entry.analysis.amenaza }} | {{ entry.analysis.analisis }}{% endif %}"
    )
    ia_run.font.size = Pt(8)
    ia_run.font.italic = True
    ia_run.font.color.rgb = RGBColor(0, 200, 200)

    _add_styled_paragraph(doc, "{% endfor %}", size=1, color=(255, 255, 255))

    endif_entries = doc.add_paragraph()
    endif_entries_run = endif_entries.add_run("{% endif %}")
    endif_entries_run.font.size = Pt(1)
    endif_entries_run.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    if_briefing = doc.add_paragraph()
    if_briefing_run = if_briefing.add_run("{% if briefing %}")
    if_briefing_run.font.size = Pt(1)
    if_briefing_run.font.color.rgb = RGBColor(255, 255, 255)

    _add_styled_paragraph(
        doc, "7. BRIEFING DE INTELIGENCIA (IA)",
        size=12, bold=True, color=(100, 255, 100),
    )
    _add_styled_paragraph(doc, "{{ briefing }}", size=10)

    endif_briefing = doc.add_paragraph()
    endif_briefing_run = endif_briefing.add_run("{% endif %}")
    endif_briefing_run.font.size = Pt(1)
    endif_briefing_run.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    hr2 = doc.add_paragraph()
    hr2_run = hr2.add_run("─" * 80)
    hr2_run.font.size = Pt(6)
    hr2_run.font.color.rgb = RGBColor(0, 180, 255)

    _add_styled_paragraph(
        doc, "COBALTO HUB v10.0 — Sistema C4I de Inteligencia OSINT",
        size=8, bold=False, color=(80, 80, 80),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_styled_paragraph(
        doc, "Documento generado automaticamente — Sin clasificar — Solo para uso institucional",
        size=7, bold=False, color=(80, 80, 80),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Plantilla generada: {OUTPUT_PATH}")


if __name__ == "__main__":
    generate_template()
